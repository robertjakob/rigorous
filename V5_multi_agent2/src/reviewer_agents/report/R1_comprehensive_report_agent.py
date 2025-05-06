from typing import Dict, Any, List
import json
import os
from datetime import datetime
from ...core.base_agent import BaseReviewerAgent
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ComprehensiveReportAgent(BaseReviewerAgent):
    """Agent responsible for generating a comprehensive report using GPT-4.1."""
    
    def __init__(self, model="gpt-4.1"):
        super().__init__(model)
        self.name = "R1_Comprehensive_Report_Agent"
        self.category = "Report Generation"
        # Initialize OpenAI
        self.client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        self.model = model
        
    def generate_comprehensive_report(self, manuscript_text: str, markdown_report: str) -> Dict[str, Any]:
        """Generates a comprehensive report using GPT-4.1."""
        try:
            # Split the markdown report into sections
            sections = self._split_markdown_into_sections(markdown_report)
            
            # Process each section separately
            analysis = {}
            for section_name, section_content in sections.items():
                # Generate prompt for this section
                prompt = f"""You are a scientific manuscript reviewer. Your task is to analyze this section of the manuscript and markdown feedback, then provide a structured JSON response with detailed feedback.

Section: {section_name}

Manuscript Text:
{manuscript_text[:1000]}...  # First 1000 chars for context

Section Feedback:
{section_content}

Structure your response as a JSON object with the following format:
{{
    "{section_name}": {{
        "score": 0-10,  // Numerical score for this section
        "summary": "A concise summary of the key points and recommendations for this section",
        "remarks": [
            {{
                "remark": "What could be an issue?",
                "original_text": "Citation of the text passage related to the issue",
                "improved_version": "New improved text",
                "explanation": "How does this new version improve the paper relative to the issue"
            }}
        ],
        "not_applicable": false  // Set to true if section doesn't apply
    }}
}}

Important:
1. Respond ONLY with the JSON object
2. Do not include any additional text or explanations
3. Ensure all text fields are properly escaped for JSON
4. Make sure the response is valid JSON that can be parsed
5. Only include genuinely helpful feedback
6. For sections that don't apply, set not_applicable to true
7. Focus on providing clear, actionable feedback without mentioning sources
8. Include numerical scores for each section
9. Extract specific feedback from the markdown report and incorporate it into your response
10. If a section has no feedback in the markdown report, provide your own assessment based on the manuscript
11. Ensure each section has at least one remark if applicable
12. Make sure the feedback is specific and actionable
13. Use the exact scores and feedback from the markdown report when available
14. Maintain the same level of detail and specificity as the markdown report
15. Include all critical remarks and improvement suggestions from the markdown report
"""
                
                # Generate content with GPT-4.1
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "You are a scientific manuscript reviewer. Your task is to analyze the manuscript and markdown feedback, then provide a structured JSON response with detailed feedback for each section."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,  # Lower temperature for more focused output
                    max_tokens=4000,  # Increased token limit
                    top_p=0.8,
                    frequency_penalty=0.0,
                    presence_penalty=0.0
                )
                
                response_text = response.choices[0].message.content.strip()
                
                # Try to parse the JSON response
                try:
                    section_analysis = json.loads(response_text)
                    analysis.update(section_analysis)
                except json.JSONDecodeError as e:
                    # If JSON parsing fails, try to extract JSON from the response
                    import re
                    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                    if json_match:
                        try:
                            section_analysis = json.loads(json_match.group())
                            analysis.update(section_analysis)
                        except json.JSONDecodeError:
                            # If still fails, try to clean the response
                            cleaned_text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_match.group())
                            section_analysis = json.loads(cleaned_text)
                            analysis.update(section_analysis)
                    else:
                        print(f"Could not parse JSON for section {section_name}: {response_text}")
                        continue
            
            # Validate the analysis structure
            required_sections = {
                "Section-Specific Agents": [f"S{i}" for i in range(1, 11)],
                "Rigor Agents": [f"R{i}" for i in range(1, 8)],
                "Writing Agents": [f"W{i}" for i in range(1, 9)]
            }
            
            for section_name, subsections in required_sections.items():
                for subsection in subsections:
                    if subsection not in analysis:
                        analysis[subsection] = {
                            "score": 0,
                            "summary": "No feedback available for this section.",
                            "remarks": [],
                            "not_applicable": True
                        }
            
            # Generate Word document report
            doc_path = self._generate_word_report(analysis, manuscript_text, markdown_report)
            
            return {
                "report_generation_score": 10,  # Assuming successful generation
                "doc_path": doc_path,
                "analysis": analysis,
                "summary": "Successfully generated comprehensive report",
                "error": False
            }
            
        except Exception as e:
            return self._generate_error_report(f"Error generating comprehensive report: {str(e)}")
    
    def _split_markdown_into_sections(self, markdown_report: str) -> Dict[str, str]:
        """Split the markdown report into sections."""
        sections = {}
        
        # Define section patterns
        section_patterns = {
            "Section-Specific Agents": {
                "S1": "## S1 - Title and Keywords",
                "S2": "## S2 - Abstract",
                "S3": "## S3 - Introduction",
                "S4": "## S4 - Literature Review",
                "S5": "## S5 - Methodology",
                "S6": "## S6 - Results",
                "S7": "## S7 - Discussion",
                "S8": "## S8 - Conclusion",
                "S9": "## S9 - References",
                "S10": "## S10 - Supplementary Materials"
            },
            "Rigor Agents": {
                "R1": "## R1 - Originality and Contribution",
                "R2": "## R2 - Impact and Significance",
                "R3": "## R3 - Ethics and Compliance",
                "R4": "## R4 - Data and Code Availability",
                "R5": "## R5 - Statistical Rigor",
                "R6": "## R6 - Technical Accuracy",
                "R7": "## R7 - Consistency"
            },
            "Writing Agents": {
                "W1": "## W1 - Language and Style",
                "W2": "## W2 - Narrative and Structure",
                "W3": "## W3 - Clarity and Conciseness",
                "W4": "## W4 - Terminology Consistency",
                "W5": "## W5 - Inclusive Language",
                "W6": "## W6 - Citation Formatting",
                "W7": "## W7 - Target Audience Alignment",
                "W8": "## W8 - Visual Presentation"
            }
        }
        
        # Split the report into sections
        lines = markdown_report.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            # Check if this line starts a new section
            found_section = False
            for category, patterns in section_patterns.items():
                for section_id, pattern in patterns.items():
                    if line.startswith(pattern):
                        # Save the previous section if it exists
                        if current_section:
                            sections[current_section] = '\n'.join(current_content)
                        # Start a new section
                        current_section = section_id
                        current_content = [line]
                        found_section = True
                        break
                if found_section:
                    break
            
            # If not a new section, add to current content
            if not found_section and current_section:
                current_content.append(line)
        
        # Save the last section
        if current_section:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _generate_word_report(self, analysis: Dict[str, Any], manuscript_text: str, markdown_report: str) -> str:
        """Generates a Word document report."""
        # Create output directory if it doesn't exist
        output_dir = "results/reports"
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = os.path.join(output_dir, f"comprehensive_report_{timestamp}.docx")
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading("Comprehensive Manuscript Review Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Add overall summary
        doc.add_heading("Overall Assessment", 1)
        summary_para = doc.add_paragraph()
        summary_para.add_run("This report provides a comprehensive review of the manuscript, analyzing its content, structure, and scientific rigor. The assessment is organized into three main categories: Section-Specific Analysis, Rigor Assessment, and Writing Quality Evaluation.")
        doc.add_paragraph()
        
        # Process each section
        sections = {
            "Section-Specific Analysis": {f"S{i}": analysis.get(f"S{i}", {}) for i in range(1, 11)},
            "Rigor Assessment": {f"R{i}": analysis.get(f"R{i}", {}) for i in range(1, 8)},
            "Writing Quality Evaluation": {f"W{i}": analysis.get(f"W{i}", {}) for i in range(1, 9)}
        }
        
        for section_name, subsections in sections.items():
            # Add section heading
            doc.add_heading(section_name, 1)
            
            # Add section introduction
            intro_text = {
                "Section-Specific Analysis": "This section provides a detailed analysis of each component of the manuscript, from title to supplementary materials.",
                "Rigor Assessment": "This section evaluates the scientific rigor, methodology, and technical aspects of the research.",
                "Writing Quality Evaluation": "This section assesses the manuscript's writing quality, clarity, and adherence to academic standards."
            }
            doc.add_paragraph(intro_text[section_name])
            doc.add_paragraph()
            
            for subsection_name, subsection_data in subsections.items():
                if subsection_data.get("not_applicable", False):
                    para = doc.add_paragraph()
                    para.add_run(f"{subsection_name}: Not applicable").bold = True
                    continue
                
                # Add subsection heading
                doc.add_heading(subsection_name, 2)
                
                # Add score with color-coded indicator
                if "score" in subsection_data:
                    score = subsection_data["score"]
                    score_para = doc.add_paragraph()
                    score_run = score_para.add_run(f"Score: {score}/10")
                    score_run.bold = True
                    # Color code: Red (<5), Yellow (5-7), Green (>7)
                    if score < 5:
                        score_run.font.color.rgb = RGBColor(192, 0, 0)  # Dark red
                    elif score < 7:
                        score_run.font.color.rgb = RGBColor(255, 192, 0)  # Orange
                    else:
                        score_run.font.color.rgb = RGBColor(0, 176, 80)  # Green
                
                # Add section summary
                if "summary" in subsection_data:
                    summary_para = doc.add_paragraph()
                    summary_para.add_run("Summary:").bold = True
                    summary_para.add_run(f"\n{subsection_data['summary']}")
                
                # Add remarks
                remarks = subsection_data.get("remarks", [])
                if remarks:
                    doc.add_heading("Detailed Feedback", 3)
                    for i, remark in enumerate(remarks, 1):
                        # Add remark number
                        remark_para = doc.add_paragraph()
                        remark_para.add_run(f"Remark {i}").bold = True
                        
                        # Add issue
                        issue_para = doc.add_paragraph()
                        issue_para.add_run("Issue: ").bold = True
                        issue_para.add_run(remark["remark"])
                        
                        # Add original text
                        original_para = doc.add_paragraph()
                        original_para.add_run("Original Text: ").bold = True
                        original_para.add_run(remark["original_text"])
                        
                        # Add improved version
                        improved_para = doc.add_paragraph()
                        improved_para.add_run("Improved Version: ").bold = True
                        improved_para.add_run(remark["improved_version"])
                        
                        # Add explanation
                        explanation_para = doc.add_paragraph()
                        explanation_para.add_run("Explanation: ").bold = True
                        explanation_para.add_run(remark["explanation"])
                        
                        doc.add_paragraph()  # Add space between remarks
                else:
                    doc.add_paragraph("No specific remarks for this section.")
                
                doc.add_paragraph()  # Add space between subsections
            
            doc.add_page_break()  # Add page break between major sections
        
        # Save document
        doc.save(doc_path)
        
        return doc_path
    
    def _generate_error_report(self, error_message: str) -> Dict[str, Any]:
        """Generates a structured error report."""
        return {
            "report_generation_score": 0,
            "doc_path": None,
            "analysis": None,
            "summary": f"Error in report generation: {error_message}",
            "error": True
        } 